# backend.py
import os
import re
import json
import time
import shutil
import subprocess
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from io import BytesIO

import httpx
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, Response

APP_DIR = Path(__file__).resolve().parent
CONTRACTS_DIR = APP_DIR / "contracts"
INCOMING_DIR = APP_DIR / "incoming"
GENERATED_DIR = APP_DIR / "generated"

INCOMING_PROPOSAL_PDF = INCOMING_DIR / "proposal.pdf"

# LLM Configuration
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "openai")  # "openai" or "ollama"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://127.0.0.1:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma3:latest")
OLLAMA_GENERATE_ENDPOINT = f"{OLLAMA_HOST}/api/generate"

MAX_PROMPT_CHARS = int(os.getenv("MAX_PROMPT_CHARS", "200000"))
OLLAMA_TIMEOUT_SECONDS = int(os.getenv("OLLAMA_TIMEOUT_SECONDS", "600"))

app = FastAPI()


def _ensure_dirs() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    INCOMING_DIR.mkdir(parents=True, exist_ok=True)
    CONTRACTS_DIR.mkdir(parents=True, exist_ok=True)


def _now_stamp() -> str:
    return time.strftime("%Y%m%d_%H%M%S")


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", errors="ignore")


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _pdftotext_available() -> bool:
    return shutil.which("pdftotext") is not None


def _pdf_to_text_pdftotext(pdf_path: Path) -> str:
    cmd = ["pdftotext", "-layout", "-nopgbrk", str(pdf_path), "-"]
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
    if p.returncode != 0:
        raise RuntimeError(f"pdftotext failed for {pdf_path.name}: {p.stderr.decode('utf-8', errors='ignore')}")
    return p.stdout.decode("utf-8", errors="ignore")


def _pdf_to_text_pypdf(pdf_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:
        raise RuntimeError(f"pypdf import failed: {e}")

    reader = PdfReader(str(pdf_path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def _pdf_to_text(pdf_path: Path) -> str:
    if _pdftotext_available():
        return _pdf_to_text_pdftotext(pdf_path)
    return _pdf_to_text_pypdf(pdf_path)


def _normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _clean_docusign_ids(text: str) -> str:
    """Remove DocuSign envelope IDs from text."""
    # Remove lines like "DocuSign Envelope ID: 15839C88-8219-4C9C-B4E1-2347F30A4DF4"
    # Also catch variations like "DocuSign Envelope ID: [TBD]" or just "DocuSign Envelope ID:"
    # Pattern matches: "DocuSign Envelope ID:" followed by UUID pattern OR [TBD] OR nothing, with optional whitespace/newlines
    text = re.sub(r'DocuSign\s+Envelope\s+ID:?\s*([A-F0-9]{8}-[A-F0-9]{4}-[A-F0-9]{4}-[A-F0-9]{4}-[A-F0-9]{12}|\[TBD\]|\[.*?\])\s*\n?', '', text, flags=re.IGNORECASE)
    # Also remove standalone "DocuSign Envelope ID:" lines (with nothing after)
    text = re.sub(r'DocuSign\s+Envelope\s+ID:?\s*\n', '', text, flags=re.IGNORECASE)
    # Remove any remaining "DocuSign" references
    text = re.sub(r'\bDocuSign\b', '', text, flags=re.IGNORECASE)
    # Also handle any trailing newlines that might be left
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _remove_page_numbers(text: str) -> str:
    """Remove page number headers/footers like 'Schedule A Page 91', 'Schedule A Page 101', etc."""
    # Remove patterns like "Schedule A Page 91", "Schedule A Page 101", etc.
    # Matches: "Schedule" followed by optional letter/number, "Page" followed by digits
    text = re.sub(r'Schedule\s+[A-Z0-9]+\s+Page\s+\d+\s*\n?', '', text, flags=re.IGNORECASE)
    # Also catch variations like "Schedule A - Page 91" or "Schedule A: Page 91"
    text = re.sub(r'Schedule\s+[A-Z0-9]+[:\-]?\s*Page\s+\d+\s*\n?', '', text, flags=re.IGNORECASE)
    # Catch standalone "Page X" patterns that might be headers/footers
    # But be careful - only remove if it's on its own line or with minimal context
    text = re.sub(r'^\s*Page\s+\d+\s*$', '', text, flags=re.IGNORECASE | re.MULTILINE)
    # Clean up extra newlines left behind
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _trim_to_specific_scope(text: str) -> str:
    """Trim output to start at 'Specific Scope' or first major section heading."""
    # Look for "Specific Scope" (case insensitive)
    scope_match = re.search(r'^\s*Specific\s+Scope', text, flags=re.IGNORECASE | re.MULTILINE)
    if scope_match:
        return text[scope_match.start():].lstrip()
    
    # If "Specific Scope" not found, look for other common section headings
    # Common patterns: numbered sections like "1.", "1.1", or major headings
    section_patterns = [
        r'^\s*\d+\.\s+[A-Z]',  # "1. SECTION NAME"
        r'^\s*[A-Z][A-Z\s]{10,}',  # All caps heading (at least 10 chars)
        r'^\s*SECTION\s+\d+',  # "SECTION 1"
        r'^\s*ARTICLE\s+\d+',  # "ARTICLE 1"
    ]
    
    for pattern in section_patterns:
        match = re.search(pattern, text, flags=re.MULTILINE)
        if match:
            # Find the start of the line containing this match
            line_start = text.rfind('\n', 0, match.start()) + 1
            return text[line_start:].lstrip()
    
    # If no section found, return as-is (better than losing content)
    return text


def _remove_company_name(text: str) -> str:
    """Remove company name (Wallcraft) and employee names for privacy. Preserves line structure."""
    # Remove employee names
    text = re.sub(r'\bDanny\s+Gates\b', '', text, flags=re.IGNORECASE)
    
    # Replace "Wallcraft Drywall, Inc." and variations
    text = re.sub(r'Wallcraft\s+Drywall,?\s*Inc\.?', '', text, flags=re.IGNORECASE)
    # Replace "Wallcraft Drywall" (without Inc.)
    text = re.sub(r'Wallcraft\s+Drywall', '', text, flags=re.IGNORECASE)
    # Replace possessive forms like "Wallcraft's"
    text = re.sub(r'Wallcraft\'?s\s+', '', text, flags=re.IGNORECASE)
    # Replace standalone "Wallcraft"
    text = re.sub(r'\bWallcraft\b', '', text, flags=re.IGNORECASE)
    
    # Clean up empty parentheses left behind
    text = re.sub(r'\(\s*\)', '', text)
    # Clean up awkward spacing patterns (but preserve newlines)
    text = re.sub(r'\s+and\s+and\s+', ' and ', text, flags=re.IGNORECASE)  # "between X and and Y"
    text = re.sub(r'\s+,\s+,', ',', text)  # Fix double commas
    # Only collapse multiple spaces on the same line (not newlines)
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space (preserves \n)
    return text


def _load_contract_text(job_name: str) -> str:
    """Load contract text from a job directory. Prioritizes PDF over TXT."""
    job_dir = CONTRACTS_DIR / job_name
    pdf_path = job_dir / "contract.pdf"
    txt_path = job_dir / "contract.txt"

    # Prioritize PDF
    if pdf_path.exists():
        txt = _pdf_to_text(pdf_path)
        txt = _normalize_text(txt)
        txt = _remove_page_numbers(txt)
        return txt

    # Fallback to TXT if PDF doesn't exist
    if txt_path.exists():
        return _normalize_text(_read_text(txt_path))

    raise FileNotFoundError(f"Missing contract for {job_name}. Expected {pdf_path} or {txt_path}")


def _load_proposal_text(job_name: str) -> str:
    """Load proposal text from a job directory."""
    job_dir = CONTRACTS_DIR / job_name
    proposal_path = job_dir / "proposal.pdf"

    if not proposal_path.exists():
        raise FileNotFoundError(f"Missing proposal for {job_name}. Expected {proposal_path}")

    txt = _pdf_to_text(proposal_path)
    txt = _normalize_text(txt)
    txt = _remove_page_numbers(txt)
    return txt


def _get_available_jobs() -> List[str]:
    """Get list of available job directories."""
    jobs = sorted([d.name for d in CONTRACTS_DIR.iterdir() if d.is_dir() and d.name.startswith("Job_")])
    return jobs


def _load_all_past_jobs() -> List[Tuple[str, str, str]]:
    """Load all past jobs (Job_001, Job_002, etc.). Returns list of (job_name, contract_text, proposal_text)."""
    jobs = _get_available_jobs()
    results = []
    
    for job_name in jobs:
        try:
            contract_text = _load_contract_text(job_name)
            proposal_text = _load_proposal_text(job_name)
            results.append((job_name, contract_text, proposal_text))
        except FileNotFoundError as e:
            # Skip jobs with missing files
            continue
    
    if not results:
        raise RuntimeError(f"No complete job data found in {CONTRACTS_DIR}")
    
    return results


def _load_new_proposal_text() -> str:
    if not INCOMING_PROPOSAL_PDF.exists():
        raise FileNotFoundError(f"Missing {INCOMING_PROPOSAL_PDF}")
    txt = _pdf_to_text(INCOMING_PROPOSAL_PDF)
    txt = _normalize_text(txt)
    txt = _remove_page_numbers(txt)
    return txt


def _build_prompt(past_jobs: List[Tuple[str, str, str]], new_proposal_text: str) -> str:
    """Build prompt using multiple past jobs to learn company style."""
    
    # Use the first past contract as the primary template
    primary_job_name, primary_contract, _ = past_jobs[0]
    
    # Build examples section from past jobs
    examples_section = "PAST JOB EXAMPLES (for style reference)\n"
    examples_section += "Study these past contracts and proposals to understand the company's style, clauses, formatting, payment terms, insurance, etc.\n\n"
    
    for job_name, contract_text, proposal_text in past_jobs:
        # Truncate if needed to fit within prompt limits
        contract_sample = contract_text[:8000] if len(contract_text) > 8000 else contract_text
        proposal_sample = proposal_text[:4000] if len(proposal_text) > 4000 else proposal_text
        
        examples_section += f"--- {job_name} ---\n"
        examples_section += f"CONTRACT:\n{contract_sample}\n\n"
        examples_section += f"PROPOSAL:\n{proposal_sample}\n\n"
    
    prompt = (
        "TASK\n"
        "You must output a full revised subcontract document that matches the company's style from the past jobs shown below.\n\n"
        "NON NEGOTIABLE RULES\n"
        "1) Output must be the entire contract text, not a summary.\n"
        "2) Match the company's exact style, clauses, formatting, payment terms, insurance, and structure from the past contracts.\n"
        "3) PRESERVE ALL LINE BREAKS, NUMBERING, AND FORMATTING from the base contract. Do not collapse everything into one line.\n"
        "4) Use the primary contract structure as the base template, but incorporate best practices from all past contracts.\n"
        "5) Only edit sections that must change to match the new proposal (scope, price, dates, subs, etc.).\n"
        "6) Do not invent anything. If something is missing, leave it unchanged or write [TBD] only in the exact place needed.\n"
        "7) Preserve numbering and lists. Do not duplicate lines.\n"
        "8) Do not add new sections like Project Overview, Key Changes, Updated Contract, Pricing Summary unless the base contract already contains them.\n"
        "9) Do NOT include DocuSign envelope IDs or any signature/tracking IDs from past contracts. Remove these completely, do not replace with [TBD].\n"
        "10) Do NOT include any company names or business entity names (like 'Wallcraft', 'Wallcraft Drywall, Inc.', etc.). Leave these fields blank.\n"
        "11) Do NOT include any employee names or personal names from past contracts. Leave these fields blank for the user to fill in.\n"
        "12) Output only the revised contract text. No commentary.\n\n"
        f"{examples_section}\n"
        "PRIMARY BASE CONTRACT (use as template)\n"
        "<<<BASE_CONTRACT\n"
        f"{primary_contract}\n"
        "BASE_CONTRACT>>>\n\n"
        "NEW PROPOSAL (update contract to match this)\n"
        "<<<NEW_PROPOSAL\n"
        f"{new_proposal_text}\n"
        "NEW_PROPOSAL>>>\n\n"
        "OUTPUT\n"
        "Return the revised contract starting from 'Specific Scope' (or the first major section heading if 'Specific Scope' is not present). Do NOT include any header information, contract numbers, or preliminary text before the main contract sections. The output must begin with the first substantive section heading.\n"
    )

    if len(prompt) > MAX_PROMPT_CHARS:
        raise RuntimeError(
            f"Prompt too large: {len(prompt)} chars. Increase MAX_PROMPT_CHARS or reduce input size."
        )
    return prompt


def _openai_generate(prompt: str, model: str, api_key: str) -> str:
    """Generate text using OpenAI API."""
    try:
        from openai import OpenAI
    except ImportError:
        raise RuntimeError("openai package not installed. Install with: pip install openai")
    
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY environment variable not set")
    
    client = OpenAI(api_key=api_key)
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a contract drafting assistant. Generate complete, professional contract documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=8000,
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        error_type = type(e).__name__
        error_msg = str(e)
        
        # Provide more specific error messages based on error type/content
        if "authentication" in error_msg.lower() or "401" in error_msg or error_type == "AuthenticationError":
            raise RuntimeError(f"OpenAI authentication error: Invalid API key. Check your OPENAI_API_KEY environment variable. Details: {error_msg}")
        elif "rate limit" in error_msg.lower() or "429" in error_msg or error_type == "RateLimitError":
            raise RuntimeError(f"OpenAI rate limit error: Too many requests. Please wait and try again. Details: {error_msg}")
        elif "context_length_exceeded" in error_msg.lower() or "token" in error_msg.lower():
            raise RuntimeError(f"OpenAI context length error: Prompt is too long. Try reducing the number of past jobs or contract length. Details: {error_msg}")
        else:
            raise RuntimeError(f"OpenAI API error ({error_type}): {error_msg}")


def _ollama_generate(prompt: str, model: str) -> str:
    """Generate text using Ollama API."""
    payload: Dict[str, Any] = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,
            "top_p": 0.9,
            "num_predict": 8000,
        },
    }

    with httpx.Client(timeout=OLLAMA_TIMEOUT_SECONDS) as client:
        r = client.post(OLLAMA_GENERATE_ENDPOINT, json=payload)

    if r.status_code != 200:
        raise RuntimeError(f"Ollama error {r.status_code}: {r.text}")

    j = r.json()
    if "response" not in j:
        raise RuntimeError(f"Unexpected Ollama payload keys: {list(j.keys())}")

    return j["response"]


def _llm_generate(prompt: str, provider: str) -> str:
    """Generate text using the specified LLM provider."""
    if provider == "openai":
        return _openai_generate(prompt, OPENAI_MODEL, OPENAI_API_KEY)
    elif provider == "ollama":
        return _ollama_generate(prompt, OLLAMA_MODEL)
    else:
        raise RuntimeError(f"Unknown LLM provider: {provider}")


def _summary_guard(final_text: str) -> None:
    bad_markers = [
        "updated contract",
        "project overview",
        "key changes",
        "pricing summary",
        "here’s a breakdown",
        "here's a breakdown",
        "implications",
    ]
    lower = final_text.lower()
    if any(m in lower for m in bad_markers):
        raise RuntimeError("Model returned summary style output. Base contract text is likely not clean enough or the model ignored rules.")


def _write_debug_files(past_jobs: List[Tuple[str, str, str]], new_proposal_text: str, prompt: str) -> None:
    """Write debug files for troubleshooting."""
    # Write combined past contracts
    past_contracts_text = "\n\n".join([f"=== {job_name} ===\n{contract}" for job_name, contract, _ in past_jobs])
    _write_text(GENERATED_DIR / "latest_debug_base_contract.txt", past_contracts_text)
    _write_text(GENERATED_DIR / "latest_debug_new_proposal.txt", new_proposal_text)
    _write_text(GENERATED_DIR / "latest_debug_prompt.txt", prompt)


def _text_to_docx(text: str, output_path: Path) -> None:
    """Convert text to Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise RuntimeError("python-docx package not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Split text into paragraphs and add to document
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
        else:
            doc.add_paragraph()  # Empty line
    
    # Save document
    doc.save(str(output_path))


@app.get("/", response_class=HTMLResponse)
def home() -> str:
    return """
<!doctype html>
<html>
  <head>
    <meta charset="utf-8"/>
    <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
    <title>Bid.AI - Contract Generator</title>
    <style>
      * { box-sizing: border-box; margin: 0; padding: 0; }
      
      :root {
        --primary: #6366f1;
        --primary-dark: #4f46e5;
        --secondary: #10b981;
        --secondary-dark: #059669;
        --bg: #0f172a;
        --bg-light: #1e293b;
        --bg-card: #1e293b;
        --text: #f1f5f9;
        --text-muted: #94a3b8;
        --border: #334155;
        --success: #22c55e;
        --error: #ef4444;
        --info: #3b82f6;
        --shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.3), 0 2px 4px -1px rgba(0, 0, 0, 0.2);
        --shadow-lg: 0 20px 25px -5px rgba(0, 0, 0, 0.3), 0 10px 10px -5px rgba(0, 0, 0, 0.2);
      }
      
      body { 
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Inter', 'Roboto', sans-serif; 
        padding: 0;
        margin: 0;
        background: linear-gradient(135deg, var(--bg) 0%, #1e293b 100%);
        color: var(--text);
        min-height: 100vh;
        line-height: 1.6;
      }
      
      .container {
        max-width: 1400px;
        margin: 0 auto;
        padding: 2rem;
      }
      
      header {
        text-align: center;
        margin-bottom: 3rem;
        padding: 2rem 0;
      }
      
      h1 {
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 0.5rem;
        letter-spacing: -0.02em;
      }
      
      .subtitle {
        color: var(--text-muted);
        font-size: 1.125rem;
        margin-top: 0.5rem;
      }
      
      .section {
        background: var(--bg-card);
        border-radius: 16px;
        padding: 2rem;
        margin-bottom: 2rem;
        box-shadow: var(--shadow);
        border: 1px solid var(--border);
        transition: transform 0.2s, box-shadow 0.2s;
      }
      
      .section:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
      }
      
      .section h2 {
        margin: 0 0 1.5rem 0;
        color: var(--text);
        font-size: 1.5rem;
        font-weight: 600;
        display: flex;
        align-items: center;
        gap: 0.75rem;
      }
      
      .section h2::before {
        content: '';
        width: 4px;
        height: 24px;
        background: linear-gradient(135deg, var(--primary) 0%, var(--secondary) 100%);
        border-radius: 2px;
      }
      
      .jobs-list {
        display: flex;
        flex-wrap: wrap;
        justify-content: center;
        gap: 2rem;
        margin-top: 1rem;
        max-width: 100%;
      }
      
      .job-card {
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.1) 0%, rgba(16, 185, 129, 0.1) 100%);
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 2rem;
        transition: all 0.3s ease;
        cursor: default;
        min-width: 450px;
        max-width: 500px;
        flex: 0 1 auto;
      }
      
      .job-card:hover {
        transform: translateY(-4px);
        border-color: var(--primary);
        box-shadow: 0 8px 16px rgba(99, 102, 241, 0.3);
      }
      
      .job-card h3 {
        margin: 0 0 1rem 0;
        font-size: 1.125rem;
        color: var(--text);
        font-weight: 600;
      }
      
      .job-card .files {
        font-size: 0.875rem;
        color: var(--text-muted);
      }
      
      .file-selection {
        display: flex;
        flex-direction: column;
        gap: 1.25rem;
        margin: 1.5rem 0;
        width: 100%;
      }
      
      .file-select-item {
        display: flex;
        align-items: center;
        justify-content: space-between;
        padding: 1.5rem 2rem;
        background: rgba(15, 23, 42, 0.3);
        border: 2px solid var(--border);
        border-radius: 10px;
        transition: all 0.3s ease;
        gap: 2rem;
        margin-bottom: 1.25rem;
        min-height: 4.5rem;
        box-sizing: border-box;
        width: 100%;
      }
      
      .file-select-item.available {
        border-color: rgba(99, 102, 241, 0.5);
        background: rgba(99, 102, 241, 0.05);
      }
      
      .file-select-item.missing {
        border-color: rgba(239, 68, 68, 0.5);
        background: rgba(239, 68, 68, 0.05);
        opacity: 0.7;
      }
      
      .file-select-item.selected {
        border-color: var(--success);
        background: rgba(34, 197, 94, 0.1);
        box-shadow: 0 0 0 3px rgba(34, 197, 94, 0.1);
      }
      
      .file-info {
        display: flex;
        align-items: center;
        gap: 1.25rem;
        flex: 1;
        min-width: 0;
        overflow: visible;
      }
      
      .file-icon {
        font-size: 1.5rem;
        flex-shrink: 0;
      }
      
      .file-name {
        color: var(--text);
        font-weight: 500;
        flex: 0 1 auto;
        white-space: nowrap;
        overflow: visible;
        text-overflow: ellipsis;
        min-width: 120px;
        margin-right: 0.5rem;
      }
      
      .file-name a {
        color: var(--text);
        text-decoration: none;
      }
      
      .file-name a:hover {
        color: var(--text);
        text-decoration: underline;
      }
      
      .file-select-item.missing .file-name {
        color: var(--text-muted);
      }
      
      .file-select-item.missing .file-name a {
        color: var(--text-muted);
      }
      
      .view-link {
        color: var(--primary);
        text-decoration: none;
        font-size: 0.875rem;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        background: rgba(99, 102, 241, 0.15);
        transition: all 0.2s;
        white-space: nowrap;
        flex-shrink: 0;
        display: inline-block;
        margin-left: 0.5rem;
      }
      
      .view-link:hover {
        background: rgba(99, 102, 241, 0.3);
        text-decoration: underline;
        transform: scale(1.05);
      }
      
      .check-btn {
        width: 3rem;
        height: 3rem;
        border-radius: 8px;
        border: 2px solid var(--border);
        background: rgba(15, 23, 42, 0.5);
        color: var(--text);
        font-size: 1.5rem;
        cursor: pointer;
        display: flex;
        align-items: center;
        justify-content: center;
        transition: all 0.2s ease;
        flex-shrink: 0;
        margin-left: 1rem;
      }
      
      .check-btn:hover:not(:disabled) {
        border-color: var(--primary);
        background: rgba(99, 102, 241, 0.1);
        transform: scale(1.1);
      }
      
      .check-btn.disabled {
        cursor: not-allowed;
        opacity: 0.5;
        border-color: var(--error);
        background: rgba(239, 68, 68, 0.1);
      }
      
      .file-select-item.selected .check-btn {
        border-color: var(--success);
        background: rgba(34, 197, 94, 0.2);
        color: var(--success);
      }
      
      .upload-status {
        font-size: 0.75rem;
        color: var(--text-muted);
        margin-top: 0.5rem;
        min-height: 1rem;
      }
      
      .upload-status.success {
        color: var(--success);
      }
      
      .upload-status.error {
        color: var(--error);
      }
      
      .proposal-info {
        background: rgba(99, 102, 241, 0.1);
        border: 1px solid var(--border);
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        display: flex;
        align-items: center;
        gap: 0.75rem;
      }
      
      .proposal-info code {
        background: rgba(99, 102, 241, 0.2);
        padding: 0.25rem 0.75rem;
        border-radius: 6px;
        font-family: 'Monaco', 'Courier New', monospace;
        color: var(--primary);
        font-size: 0.875rem;
      }
      
      button {
        padding: 0.875rem 1.75rem;
        font-size: 1rem;
        background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
        color: white;
        border: none;
        border-radius: 10px;
        cursor: pointer;
        font-weight: 600;
        margin-right: 0.75rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(99, 102, 241, 0.3);
        position: relative;
        overflow: hidden;
      }
      
      button::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 50%;
        width: 0;
        height: 0;
        border-radius: 50%;
        background: rgba(255, 255, 255, 0.3);
        transform: translate(-50%, -50%);
        transition: width 0.6s, height 0.6s;
      }
      
      button:hover::before {
        width: 300px;
        height: 300px;
      }
      
      button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(99, 102, 241, 0.4);
      }
      
      button:active {
        transform: translateY(0);
      }
      
      button:disabled {
        background: #475569;
        cursor: not-allowed;
        transform: none;
        box-shadow: none;
      }
      
      button.secondary {
        background: linear-gradient(135deg, var(--secondary) 0%, var(--secondary-dark) 100%);
        box-shadow: 0 4px 6px rgba(16, 185, 129, 0.3);
      }
      
      button.secondary:hover:not(:disabled) {
        box-shadow: 0 6px 12px rgba(16, 185, 129, 0.4);
      }
      
      textarea {
        width: 100%;
        min-height: 500px;
        font-family: 'Monaco', 'Courier New', monospace;
        font-size: 0.875rem;
        padding: 1.5rem;
        background: rgba(15, 23, 42, 0.5);
        border: 1px solid var(--border);
        border-radius: 12px;
        color: var(--text);
        line-height: 1.8;
        resize: vertical;
        transition: all 0.3s ease;
      }
      
      textarea:focus {
        outline: none;
        border-color: var(--primary);
        box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.2);
      }
      
      textarea::placeholder {
        color: var(--text-muted);
      }
      
      .status {
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        display: none;
        border-left: 4px solid;
        animation: slideIn 0.3s ease;
      }
      
      @keyframes slideIn {
        from {
          opacity: 0;
          transform: translateX(-20px);
        }
        to {
          opacity: 1;
          transform: translateX(0);
        }
      }
      
      .status.info {
        background: rgba(59, 130, 246, 0.1);
        color: var(--info);
        border-color: var(--info);
      }
      
      .status.error {
        background: rgba(239, 68, 68, 0.1);
        color: var(--error);
        border-color: var(--error);
      }
      
      .status.success {
        background: rgba(34, 197, 94, 0.1);
        color: var(--success);
        border-color: var(--success);
      }
      
      .buttons-group {
        margin-top: 1.5rem;
        display: flex;
        gap: 1rem;
        flex-wrap: wrap;
      }
      
      .loading {
        display: inline-block;
        width: 16px;
        height: 16px;
        border: 2px solid rgba(255, 255, 255, 0.3);
        border-top: 2px solid white;
        border-radius: 50%;
        animation: spin 0.8s linear infinite;
        margin-left: 8px;
        vertical-align: middle;
      }
      
      @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
      }
      
      @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.5; }
      }
      
      #jobs-container:empty::after {
        content: 'Loading jobs...';
        color: var(--text-muted);
        animation: pulse 2s ease-in-out infinite;
      }
      
      @media (max-width: 768px) {
        .container {
          padding: 1rem;
        }
        
        h1 {
          font-size: 2rem;
        }
        
        .jobs-list {
          flex-direction: column;
          align-items: center;
        }
        
        .job-card {
          min-width: 100%;
          max-width: 100%;
        }
        
        .section {
          padding: 1.5rem;
        }
      }
    </style>
  </head>
  <body>
    <div class="container">
      <header>
        <h1>⚡ Bid.AI</h1>
        <p class="subtitle">AI-Powered Contract Generator</p>
        <p class="subtitle" style="font-size: 0.875rem; margin-top: 0.25rem;">
          Generate professional contracts using past job examples and new proposals
        </p>
      </header>
      
      <div class="section">
        <h2>ℹ️ How this works</h2>
        <div style="color: var(--text); line-height: 1.8; font-size: 1rem;">
          <ol style="margin-left: 1.5rem; padding-left: 0.5rem;">
            <li style="margin-bottom: 1rem;"><strong>Upload historical proposal and contract</strong><br>
            <span style="color: var(--text-muted);">Select contract and proposal PDFs from past jobs. These show your company's standard terms and style.</span></li>
            
            <li style="margin-bottom: 1rem;"><strong>Upload new proposal</strong><br>
            <span style="color: var(--text-muted);">Upload the proposal PDF for the job you want to create a contract for.</span></li>
            
            <li style="margin-bottom: 1rem;"><strong>System identifies deltas</strong><br>
            <span style="color: var(--text-muted);">The system compares the new proposal to your past contracts to find what's different.</span></li>
            
            <li style="margin-bottom: 1rem;"><strong>System edits only impacted sections</strong><br>
            <span style="color: var(--text-muted);">Only the sections that need to change are updated. Everything else stays the same.</span></li>
            
            <li style="margin-bottom: 1rem;"><strong>User reviews and finalizes</strong><br>
            <span style="color: var(--text-muted);">Review the generated contract, make any adjustments, and download when ready.</span></li>
          </ol>
        </div>
      </div>
      
      <div class="section">
        <h2>📁 Upload Job Files</h2>
        <p style="color: var(--text-muted); margin-bottom: 1.5rem;">
          Drag and drop contract and proposal files for each job. Files are required before generating contracts.
        </p>
        <div id="jobs-container" class="jobs-list"></div>
      </div>
      
      <div class="section">
        <h2>📋 New Job Proposal</h2>
        <p style="color: var(--text-muted); margin-bottom: 1rem;">
          Upload the proposal PDF for the new job you want to generate a contract for.
        </p>
        <div id="incoming-proposal-container">
          <div class="file-selection">
            <div class="file-select-item" id="incoming-file-item" data-type="incoming">
              <div class="file-info">
                <span class="file-icon">📄</span>
                <span class="file-name" id="incoming-file-link">proposal.pdf (not uploaded)</span>
              </div>
              <button class="check-btn disabled" id="incoming-check-btn" disabled data-type="incoming">
                ✗
              </button>
              <input type="file" id="file-input-incoming" 
                     style="display: none;" 
                     accept=".pdf">
            </div>
          </div>
        </div>
      </div>
      
      <div class="section">
        <h2>✨ Generate Contract</h2>
        <button id="gen-btn" disabled>🚀 Generate Draft Contract</button>
        <div id="status" class="status"></div>
      </div>
      
      <div class="section" id="change-summary-section" style="display: none;">
        <h2>📊 Change Summary</h2>
        <div id="change-summary" style="background: rgba(15, 23, 42, 0.5); border: 1px solid var(--border); border-radius: 12px; padding: 1.5rem; color: var(--text); line-height: 1.8; font-size: 1rem;">
          <div id="change-summary-content">
            <!-- Summary will be populated here -->
          </div>
        </div>
      </div>
      
      <div class="section">
        <h2>📝 Generated Contract</h2>
        <textarea id="contract-text" placeholder="Your generated contract will appear here..."></textarea>
        <div class="buttons-group">
          <button id="download-btn" class="secondary" disabled>📥 Download as Word (.docx)</button>
        </div>
      </div>
    
    <script>
      const jobsContainer = document.getElementById('jobs-container');
      const genBtn = document.getElementById('gen-btn');
      const downloadBtn = document.getElementById('download-btn');
      const contractText = document.getElementById('contract-text');
      const status = document.getElementById('status');
      const changeSummarySection = document.getElementById('change-summary-section');
      const changeSummaryContent = document.getElementById('change-summary-content');
      
      function showStatus(message, type) {
        status.textContent = message;
        status.className = 'status ' + type;
        status.style.display = 'block';
      }
      
      function hideStatus() {
        status.style.display = 'none';
      }
      
      // Load jobs on page load
      async function loadJobs() {
        try {
          const res = await fetch('/api/jobs');
          const data = await res.json();
          const jobs = data.jobs || [];
          
          if (jobs.length === 0) {
            jobsContainer.innerHTML = '<p style="color: var(--text-muted);">No jobs found. Add Job_001 and Job_002 folders in contracts/</p>';
            return;
          }
          
          jobsContainer.innerHTML = jobs.map(job => {
            const contractExists = job.has_contract_pdf;
            const proposalExists = job.has_proposal_pdf;
            
            return `
            <div class="job-card">
              <h3>${job.name}</h3>
              <div class="file-selection">
                <div class="file-select-item ${contractExists ? 'available' : 'missing'}" 
                     data-job="${job.name}" 
                     data-type="contract"
                     data-exists="${contractExists}">
                  <div class="file-info">
                    <span class="file-icon">📄</span>
                    <span class="file-name">contract.pdf</span>
                    ${contractExists ? `<a href="/api/view/${job.name}/contract_pdf" target="_blank" class="view-link">View</a>` : ''}
                  </div>
                  <button class="check-btn ${contractExists ? '' : 'disabled'}" 
                          ${contractExists ? '' : 'disabled'}
                          data-job="${job.name}" 
                          data-type="contract">
                    ${contractExists ? '☐' : '✗'}
                  </button>
                  <input type="file" id="file-input-${job.name}-contract" 
                         style="display: none;" 
                         accept=".pdf">
                </div>
                
                <div class="file-select-item ${proposalExists ? 'available' : 'missing'}" 
                     data-job="${job.name}" 
                     data-type="proposal"
                     data-exists="${proposalExists}">
                  <div class="file-info">
                    <span class="file-icon">📋</span>
                    <span class="file-name">proposal.pdf</span>
                    ${proposalExists ? `<a href="/api/view/${job.name}/proposal_pdf" target="_blank" class="view-link">View</a>` : ''}
                  </div>
                  <button class="check-btn ${proposalExists ? '' : 'disabled'}" 
                          ${proposalExists ? '' : 'disabled'}
                          data-job="${job.name}" 
                          data-type="proposal">
                    ${proposalExists ? '☐' : '✗'}
                  </button>
                  <input type="file" id="file-input-${job.name}-proposal" 
                         style="display: none;" 
                         accept=".pdf">
                </div>
              </div>
            </div>
          `;
          }).join('');
          
          // Setup file selection handlers
          setupFileSelection();
        } catch (e) {
          jobsContainer.innerHTML = '<p style="color: var(--error);">Error loading jobs: ' + e.message + '</p>';
        }
      }
      
      // Generate change summary based on contract text
      function generateChangeSummary(contractText) {
        const summary = [];
        
        // Analyze contract text for common changes
        const text = contractText.toLowerCase();
        
        // Check for pricing/schedule sections
        if (text.includes('schedule') || text.includes('pricing') || text.includes('amount') || text.includes('$')) {
          summary.push('Pricing updated');
        }
        
        // Check for alternates
        if (text.includes('alternate') || text.includes('addendum')) {
          summary.push('Alternates added');
        }
        
        // Check for scope changes
        if (text.includes('specific scope') || text.includes('work scope') || text.includes('project scope')) {
          summary.push('Sections modified');
        }
        
        // Check for exclusions
        if (text.includes('exclusion') || text.includes('not included')) {
          summary.push('Exclusions unchanged');
        } else {
          // If no exclusions mentioned, they might be unchanged
          summary.push('Exclusions unchanged');
        }
        
        // Always show sections modified if we have content
        if (contractText.length > 100) {
          if (!summary.includes('Sections modified')) {
            summary.unshift('Sections modified');
          }
        }
        
        // Display summary
        if (summary.length > 0) {
          changeSummaryContent.innerHTML = '<ul style="margin: 0; padding-left: 1.5rem; list-style-type: disc;">' +
            summary.map(item => `<li style="margin-bottom: 0.5rem;">${item}</li>`).join('') +
            '</ul>';
          changeSummarySection.style.display = 'block';
        } else {
          changeSummarySection.style.display = 'none';
        }
      }
      
      // Check if required files are selected
      function checkRequiredFiles() {
        const jobs = Array.from(document.querySelectorAll('.job-card'));
        let hasRequiredFiles = true;
        let missingFiles = [];
        
        // Check each job has both contract and proposal selected
        jobs.forEach(jobCard => {
          const contractItem = jobCard.querySelector('.file-select-item[data-type="contract"]');
          const proposalItem = jobCard.querySelector('.file-select-item[data-type="proposal"]');
          
          if (!contractItem || !contractItem.classList.contains('selected')) {
            hasRequiredFiles = false;
            const jobName = contractItem?.dataset.job || 'Unknown';
            if (!missingFiles.includes(`${jobName} - contract.pdf`)) {
              missingFiles.push(`${jobName} - contract.pdf`);
            }
          }
          
          if (!proposalItem || !proposalItem.classList.contains('selected')) {
            hasRequiredFiles = false;
            const jobName = proposalItem?.dataset.job || 'Unknown';
            if (!missingFiles.includes(`${jobName} - proposal.pdf`)) {
              missingFiles.push(`${jobName} - proposal.pdf`);
            }
          }
        });
        
        // Check incoming proposal
        const incomingCheckBtn = document.getElementById('incoming-check-btn');
        const incomingItem = document.getElementById('incoming-file-item');
        const hasIncoming = incomingCheckBtn && incomingCheckBtn.textContent === '✓';
        
        if (!hasIncoming) {
          hasRequiredFiles = false;
          missingFiles.push('New proposal PDF');
        }
        
        // Check job files
        jobs.forEach(card => {
          const jobName = card.querySelector('h3').textContent;
          // Check if both contract and proposal are selected
          const contractItem = card.querySelector('.file-select-item[data-type="contract"]');
          const proposalItem = card.querySelector('.file-select-item[data-type="proposal"]');
          
          const hasContract = contractItem && contractItem.classList.contains('selected');
          const hasProposal = proposalItem && proposalItem.classList.contains('selected');
          
          if (!hasContract || !hasProposal) {
            hasRequiredFiles = false;
            missingFiles.push(`${jobName}: ${!hasContract ? 'contract.pdf' : ''} ${!hasProposal ? 'proposal.pdf' : ''}`.trim());
          }
        });
        
        genBtn.disabled = !hasRequiredFiles;
        if (!hasRequiredFiles) {
          genBtn.title = 'Please upload all required files (contract.pdf and proposal.pdf for each job, plus new proposal PDF)';
        } else {
          genBtn.title = '';
        }
        
        return hasRequiredFiles;
      }
      
      // Generate contract
      genBtn.addEventListener('click', async () => {
        if (!checkRequiredFiles()) {
          showStatus('❌ Please upload all required files (contract and proposal) for each job before generating.', 'error');
          return;
        }
        
        genBtn.disabled = true;
        genBtn.innerHTML = '🔄 Generating<span class="loading"></span>';
        hideStatus();
        contractText.value = '';
        downloadBtn.disabled = true;
        changeSummarySection.style.display = 'none';
        
        try {
          showStatus('⚡ Generating contract... This may take a minute.', 'info');
          const res = await fetch('/generate', { method: 'POST' });
          
          if (!res.ok) {
            const errorData = await res.json().catch(() => ({ detail: res.statusText }));
            throw new Error(errorData.detail || `HTTP ${res.status}`);
          }
          
          const data = await res.json();
          
          if (data.contract_text) {
            contractText.value = data.contract_text;
            downloadBtn.disabled = false;
            showStatus('✅ Contract generated successfully!', 'success');
            
            // Generate and show change summary
            generateChangeSummary(data.contract_text);
            
            // Smooth scroll to change summary
            changeSummarySection.scrollIntoView({ behavior: 'smooth', block: 'start' });
          } else {
            showStatus('❌ Generated but no contract text returned.', 'error');
          }
        } catch (e) {
          showStatus('❌ Error: ' + e.message, 'error');
          changeSummarySection.style.display = 'none';
        } finally {
          genBtn.disabled = false;
          genBtn.innerHTML = '🚀 Generate Draft Contract';
          checkRequiredFiles(); // Re-check after generation
        }
      });
      
      // Download Word document
      downloadBtn.addEventListener('click', () => {
        window.location.href = '/api/download/docx';
      });
      
      // Setup file selection system
      function setupFileSelection() {
        // Handle checkmark clicks for job files (including disabled ones)
        document.querySelectorAll('.file-select-item .check-btn').forEach(btn => {
          // Skip incoming proposal button (handled separately)
          if (btn.id === 'incoming-check-btn') return;
          
          // Remove old listeners by cloning
          const newBtn = btn.cloneNode(true);
          btn.parentNode.replaceChild(newBtn, btn);
          
          newBtn.addEventListener('click', (e) => {
            e.stopPropagation();
            const item = newBtn.closest('.file-select-item');
            const jobName = item.dataset.job;
            const fileType = item.dataset.type;
            const exists = item.dataset.exists === 'true';
            
            // If disabled or file doesn't exist, open file picker
            if (newBtn.disabled || !exists) {
              const input = item.querySelector('input[type="file"]');
              if (input) input.click();
              return;
            }
            
            // Toggle selection
            if (item.classList.contains('selected')) {
              item.classList.remove('selected');
              newBtn.textContent = '☐';
            } else {
              item.classList.add('selected');
              newBtn.textContent = '✓';
            }
            
            checkRequiredFiles();
          });
        });
        
        // Handle file input changes (for uploading new files)
        document.querySelectorAll('.file-select-item input[type="file"]').forEach(input => {
          // Remove old listeners
          const newInput = input.cloneNode(true);
          input.parentNode.replaceChild(newInput, input);
          
          newInput.addEventListener('change', async (e) => {
            const file = e.target.files[0];
            if (!file) return;
            
            const item = newInput.closest('.file-select-item');
            const jobName = item.dataset.job;
            const fileType = item.dataset.type;
            
            if (fileType === 'incoming') {
              await uploadIncomingProposal(file);
            } else if (jobName && fileType) {
              await uploadFile(file, jobName, fileType);
            }
          });
        });
        
        // Handle incoming proposal checkmark (always set up, even if disabled)
        const incomingCheckBtn = document.getElementById('incoming-check-btn');
        if (incomingCheckBtn) {
          // Remove old listeners by cloning
          const newIncomingBtn = incomingCheckBtn.cloneNode(true);
          incomingCheckBtn.parentNode.replaceChild(newIncomingBtn, incomingCheckBtn);
          
          newIncomingBtn.addEventListener('click', (e) => {
            e.stopPropagation();
            
            // If disabled, don't do anything (file doesn't exist)
            if (newIncomingBtn.disabled) {
              const input = document.getElementById('file-input-incoming');
              if (input) input.click();
              return;
            }
            
            const item = document.getElementById('incoming-file-item');
            if (!item) return;
            
            // Toggle selection
            if (item.classList.contains('selected')) {
              item.classList.remove('selected');
              newIncomingBtn.textContent = '☐';
            } else {
              item.classList.add('selected');
              newIncomingBtn.textContent = '✓';
            }
            
            checkRequiredFiles();
          });
        }
      }
      
      // Function to upload incoming proposal
      async function uploadIncomingProposal(file) {
        const statusEl = document.getElementById('status-incoming');
        if (!statusEl) {
          console.error('Status element not found for incoming proposal');
          return;
        }
        
        // Validate file type
        const fileExt = file.name.toLowerCase().split('.').pop();
        if (fileExt !== 'pdf') {
          statusEl.textContent = '✗ Only .pdf files allowed';
          statusEl.className = 'upload-status error';
          return;
        }
        
        const dropZone = document.getElementById('incoming-drop-zone');
        
        // Show uploading state
        statusEl.textContent = 'Uploading...';
        statusEl.className = 'upload-status';
        if (dropZone) {
          dropZone.classList.remove('uploaded');
        }
        
        const formData = new FormData();
        formData.append('file', file);
        
        try {
          const res = await fetch('/api/upload/incoming', {
            method: 'POST',
            body: formData
          });
          
          const data = await res.json();
          
          if (res.ok) {
            // Show success with filename
            statusEl.innerHTML = `<span class="uploaded-filename">✓ ${file.name}</span>`;
            statusEl.className = 'upload-status success';
            
            // Turn drop zone green
            if (dropZone) {
              dropZone.classList.add('uploaded');
            }
            
            // Update the file link
            const fileLink = document.getElementById('incoming-file-link');
            if (fileLink) {
              fileLink.innerHTML = '<a href="/api/view/incoming" target="_blank">proposal.pdf</a>';
              fileLink.parentElement.classList.remove('missing');
            }
            // Re-check required files
            setTimeout(() => checkRequiredFiles(), 500);
          } else {
            throw new Error(data.detail || 'Upload failed');
          }
        } catch (e) {
          console.error('Upload error:', e);
          statusEl.textContent = '✗ ' + (e.message || 'Upload failed');
          statusEl.className = 'upload-status error';
          if (dropZone) {
            dropZone.classList.remove('uploaded');
          }
        }
      }
      
      // Function to check incoming proposal status
      async function checkIncomingProposal() {
        try {
          const res = await fetch('/health');
          const data = await res.json();
          const exists = data.detail?.incoming_proposal_exists || false;
          
          const fileLink = document.getElementById('incoming-file-link');
          const fileItem = document.getElementById('incoming-file-item');
          const checkBtn = document.getElementById('incoming-check-btn');
          
          if (fileLink && fileItem && checkBtn) {
            if (exists) {
              fileLink.innerHTML = '<a href="/api/view/incoming" target="_blank">proposal.pdf</a>';
              fileItem.classList.add('available');
              fileItem.dataset.exists = 'true';
              checkBtn.textContent = '☐';
              checkBtn.disabled = false;
              checkBtn.classList.remove('disabled');
              
              // Re-setup the button handler after enabling it
              setTimeout(() => setupFileSelection(), 100);
            } else {
              fileLink.textContent = 'proposal.pdf (not uploaded)';
              fileItem.classList.remove('available');
              fileItem.dataset.exists = 'false';
              checkBtn.textContent = '✗';
              checkBtn.disabled = true;
              checkBtn.classList.add('disabled');
            }
          }
        } catch (e) {
          console.error('Error checking incoming proposal:', e);
        }
      }
      
      async function uploadFile(file, jobName, fileType) {
        // Validate file type
        const fileExt = file.name.toLowerCase().split('.').pop();
        if (fileExt !== 'pdf') {
          alert('Only .pdf files are allowed');
          return;
        }
        
        const fileItem = document.querySelector(`.file-select-item[data-job="${jobName}"][data-type="${fileType}"]`);
        const checkBtn = fileItem?.querySelector('.check-btn');
        
        // Show uploading state
        if (checkBtn) {
          checkBtn.textContent = '⏳';
          checkBtn.disabled = true;
        }
        
        const formData = new FormData();
        formData.append('file', file);
        formData.append('job_name', jobName);
        formData.append('file_type', fileType);
        
        try {
          const res = await fetch('/api/upload', {
            method: 'POST',
            body: formData
          });
          
          const data = await res.json();
          
          if (res.ok) {
            // Reload jobs to update file status
            setTimeout(() => {
              loadJobs().then(() => {
                checkRequiredFiles();
                checkIncomingProposal();
                setupFileSelection();
              });
            }, 500);
          } else {
            throw new Error(data.detail || 'Upload failed');
          }
        } catch (e) {
          console.error('Upload error:', e);
          if (checkBtn) {
            checkBtn.textContent = '✗';
            checkBtn.disabled = false;
          }
          alert('Upload failed: ' + (e.message || 'Unknown error'));
        }
      }
      
      // Load jobs and check incoming proposal on page load
      loadJobs().then(() => {
        checkRequiredFiles();
        checkIncomingProposal();
        setupFileSelection();
      });
      
      // Also setup file selection immediately (for incoming proposal in static HTML)
      setTimeout(() => {
        setupFileSelection();
      }, 100);
    </script>
  </body>
</html>
""".strip()


@app.get("/api/jobs")
def get_jobs() -> JSONResponse:
    """Get list of available past jobs."""
    _ensure_dirs()
    jobs = _get_available_jobs()
    
    job_details = []
    for job_name in jobs:
        job_dir = CONTRACTS_DIR / job_name
        job_details.append({
            "name": job_name,
            "has_contract_pdf": (job_dir / "contract.pdf").exists(),
            "has_proposal_pdf": (job_dir / "proposal.pdf").exists(),
        })
    
    return JSONResponse({"jobs": job_details})


@app.post("/api/upload")
async def upload_file(
    job_name: str = Form(...),
    file_type: str = Form(...),  # "contract" or "proposal"
    file: UploadFile = File(...)
) -> JSONResponse:
    """Upload a file to a job directory."""
    _ensure_dirs()
    
    # Validate job name
    if not job_name.startswith("Job_"):
        raise HTTPException(status_code=400, detail="Job name must start with 'Job_'")
    
    # Validate file type
    if file_type not in ["contract", "proposal"]:
        raise HTTPException(status_code=400, detail="file_type must be 'contract' or 'proposal'")
    
    # Determine file extension
    original_filename = file.filename or "file"
    file_ext = Path(original_filename).suffix.lower()
    
    # Determine target filename
    if file_type == "contract":
        if file_ext == ".pdf":
            target_filename = "contract.pdf"
        else:
            raise HTTPException(status_code=400, detail="Contract files must be .pdf")
    else:  # proposal
        if file_ext == ".pdf":
            target_filename = "proposal.pdf"
        else:
            raise HTTPException(status_code=400, detail="Proposal files must be .pdf")
    
    # Create job directory if it doesn't exist
    job_dir = CONTRACTS_DIR / job_name
    job_dir.mkdir(parents=True, exist_ok=True)
    
    # Save file
    target_path = job_dir / target_filename
    try:
        content = await file.read()
        target_path.write_bytes(content)
        return JSONResponse({
            "success": True,
            "message": f"File uploaded successfully",
            "job_name": job_name,
            "file_type": file_type,
            "filename": target_filename,
            "path": str(target_path)
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")


@app.post("/api/upload/incoming")
async def upload_incoming_proposal(file: UploadFile = File(...)) -> JSONResponse:
    """Upload the incoming proposal PDF."""
    _ensure_dirs()
    
    # Validate file extension
    original_filename = file.filename or "file"
    file_ext = Path(original_filename).suffix.lower()
    
    if file_ext != ".pdf":
        raise HTTPException(status_code=400, detail="Incoming proposal must be a .pdf file")
    
    # Save file
    target_path = INCOMING_PROPOSAL_PDF
    try:
        content = await file.read()
        target_path.write_bytes(content)
        return JSONResponse({
            "success": True,
            "message": "Incoming proposal uploaded successfully",
            "filename": "proposal.pdf",
            "path": str(target_path)
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")


@app.get("/api/view/incoming")
def view_incoming_proposal() -> Response:
    """View the incoming proposal PDF."""
    _ensure_dirs()
    
    if not INCOMING_PROPOSAL_PDF.exists():
        raise HTTPException(status_code=404, detail="Incoming proposal not found")
    
    return FileResponse(
        path=str(INCOMING_PROPOSAL_PDF),
        filename="proposal.pdf",
        media_type="application/pdf"
    )


@app.get("/api/view/{job_name}/{file_type}")
def view_file(job_name: str, file_type: str) -> Response:
    """View a file from a job directory."""
    _ensure_dirs()
    
    # Validate job name
    if not job_name.startswith("Job_"):
        raise HTTPException(status_code=400, detail="Invalid job name")
    
    # Determine filename
    if file_type == "contract_txt":
        filename = "contract.txt"
        media_type = "text/plain"
    elif file_type == "contract_pdf":
        filename = "contract.pdf"
        media_type = "application/pdf"
    elif file_type == "proposal_pdf":
        filename = "proposal.pdf"
        media_type = "application/pdf"
    else:
        raise HTTPException(status_code=400, detail="Invalid file type")
    
    file_path = CONTRACTS_DIR / job_name / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    # For text files, return as HTML with syntax highlighting
    if file_type == "contract_txt":
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{job_name} - {filename}</title>
            <style>
                body {{
                    font-family: 'Monaco', 'Courier New', monospace;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 2rem;
                    background: #0f172a;
                    color: #f1f5f9;
                    line-height: 1.6;
                }}
                pre {{
                    background: #1e293b;
                    padding: 1.5rem;
                    border-radius: 8px;
                    border: 1px solid #334155;
                    overflow-x: auto;
                    white-space: pre-wrap;
                    word-wrap: break-word;
                }}
                .back-btn {{
                    display: inline-block;
                    margin-bottom: 1rem;
                    padding: 0.5rem 1rem;
                    background: #6366f1;
                    color: white;
                    text-decoration: none;
                    border-radius: 6px;
                }}
            </style>
        </head>
        <body>
            <a href="javascript:history.back()" class="back-btn">← Back</a>
            <h1>{job_name} - {filename}</h1>
            <pre>{content}</pre>
        </body>
        </html>
        """
        return HTMLResponse(content=html_content)
    
    # For PDF files, return as file download/view
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type=media_type
    )


@app.get("/health")
def health() -> JSONResponse:
    _ensure_dirs()
    ok = True
    jobs = _get_available_jobs()
    
    detail = {
        "llm_provider": LLM_PROVIDER,
        "incoming_proposal_exists": INCOMING_PROPOSAL_PDF.exists(),
        "available_jobs": jobs,
        "openai_api_key_set": bool(OPENAI_API_KEY),
    }
    
    if LLM_PROVIDER == "openai":
        detail["openai_model"] = OPENAI_MODEL
    else:
        detail["ollama_host"] = OLLAMA_HOST
        detail["ollama_model"] = OLLAMA_MODEL
        try:
            with httpx.Client(timeout=10) as client:
                tags = client.get(f"{OLLAMA_HOST}/api/tags")
            detail["ollama_reachable"] = tags.status_code == 200
        except Exception as e:
            detail["ollama_reachable"] = False
            detail["ollama_error"] = str(e)
            ok = False

    return JSONResponse({"ok": ok, "detail": detail})


@app.post("/generate")
def generate() -> JSONResponse:
    _ensure_dirs()

    try:
        # Load all past jobs (Job_001, Job_002, etc.)
        past_jobs = _load_all_past_jobs()
        
        # Load new proposal
        new_proposal = _load_new_proposal_text()
        
        # Build prompt using all past jobs
        prompt = _build_prompt(past_jobs, new_proposal)
        
        # Write debug files
        _write_debug_files(past_jobs, new_proposal, prompt)
        
        # Generate contract using LLM
        final_text = _llm_generate(prompt, LLM_PROVIDER)
        final_text = _normalize_text(final_text)
        final_text = _trim_to_specific_scope(final_text)  # Trim to start at Specific Scope
        final_text = _remove_page_numbers(final_text)
        final_text = _clean_docusign_ids(final_text)
        final_text = _remove_company_name(final_text)
        final_text = _normalize_text(final_text)  # Re-normalize after cleaning
        _summary_guard(final_text)

        # Save text file
        stamp = _now_stamp()
        out_path_txt = GENERATED_DIR / f"{stamp}_Job_003_contract.txt"
        _write_text(out_path_txt, final_text)
        _write_text(GENERATED_DIR / "latest.txt", final_text)
        
        # Save Word document
        out_path_docx = GENERATED_DIR / f"{stamp}_Job_003_contract.docx"
        _text_to_docx(final_text, out_path_docx)

        return JSONResponse({
            "saved_path": str(out_path_txt),
            "docx_path": str(out_path_docx),
            "contract_text": final_text
        })

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/download/docx")
def download_docx() -> FileResponse:
    """Download the latest generated contract as Word document."""
    _ensure_dirs()
    
    # Find the latest .docx file
    docx_files = sorted(GENERATED_DIR.glob("*_contract.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    
    if not docx_files:
        raise HTTPException(status_code=404, detail="No contract document found. Generate a contract first.")
    
    latest_docx = docx_files[0]
    return FileResponse(
        path=str(latest_docx),
        filename=latest_docx.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )